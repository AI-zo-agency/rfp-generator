"""The Submission Readiness Report — a standalone internal DOCX.

Deliberately NOT an appendix to the proposal. An internal readiness score and a list of
one's own gaps must never reach an evaluator, and the safest guarantee of that is that
they never share a file.

Rows are written for the person doing the work. A blocker names the owner, where it
lives, the RFP clause requiring it, and what happens if it ships without it — not
"Missing bid bond".
"""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor

from app.models.proposal import ManualFillFlag
from app.services.proposal_readiness import CriterionScore, ReadinessResult

logger = logging.getLogger(__name__)

_RED = RGBColor(0xB4, 0x1C, 0x1C)
_AMBER = RGBColor(0x9A, 0x5B, 0x00)
_GREY = RGBColor(0x55, 0x55, 0x55)


def build_readiness_report_filename(*, rfp_title: str) -> str:
    safe = "".join(c for c in (rfp_title or "Proposal") if c.isalnum() or c in " -_").strip()
    return f"{safe or 'Proposal'} - Submission Readiness Report.docx"


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.space_before = Pt(14)


def _kv(doc: Document, label: str, value: str, *, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    lab = p.add_run(f"{label}: ")
    lab.bold = True
    val = p.add_run(value)
    if color is not None:
        val.font.color.rgb = color


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _GREY


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value or "—"


def _blocker_rows(flags: list[ManualFillFlag]) -> list[list[str]]:
    return [
        [
            f.tag,
            f.owner or "Unassigned",
            f.section_title or f.section_id,
            # The quoted clause is the whole point of the row: it is why this is a
            # blocker rather than an assertion that it is one.
            f'"{f.rfp_evidence}"' if f.rfp_evidence else "(no clause cited)",
            f.if_skipped or "Bid may be rejected.",
        ]
        for f in flags
    ]


def build_readiness_report_docx_bytes(
    *,
    rfp_title: str,
    readiness: ReadinessResult,
    flags: list[ManualFillFlag],
    scores: list[CriterionScore] | None = None,
    changes: list[str] | None = None,
    unverified_claims: list[str] | None = None,
    unfixed: list[str] | None = None,
) -> bytes:
    """Render the report. Safe on empty inputs — a clean draft still gets a document."""
    doc = Document()

    title = doc.add_paragraph()
    trun = title.add_run("Submission Readiness Report")
    trun.bold = True
    trun.font.size = Pt(20)
    _note(doc, f"{rfp_title} · generated {datetime.now(timezone.utc):%Y-%m-%d %H:%M UTC}")
    _note(doc, "Internal document. Do not submit with the proposal.")

    # 1 — Verdict, readable in five seconds.
    _heading(doc, "1. Verdict")
    verdict_color = _RED if readiness.open_disqualifying else (
        None if readiness.ready else _AMBER
    )
    # "not measured" is not "0%" — say which one this is.
    _kv(
        doc,
        "Readiness",
        f"{readiness.score}%" if readiness.measured else "not measured",
        color=verdict_color,
    )
    _kv(doc, "Confidence", f"{readiness.confidence} — {readiness.confidence_note}")
    _kv(doc, "Verdict", readiness.verdict, color=verdict_color)

    # 2 — Blockers first: these decide whether anything else matters.
    blockers = [f for f in flags if f.criticality == "disqualifying"]
    _heading(doc, f"2. Blockers ({len(blockers)})")
    if blockers:
        _note(doc, "Each of these can render the bid non-responsive on its own.")
        _table(
            doc,
            ["What's needed", "Owner", "Where", "Why (RFP clause)", "If skipped"],
            _blocker_rows(blockers),
        )
    else:
        doc.add_paragraph("None. No disqualifying items are open.")

    # 3 — Scored gaps, ordered so effort goes where the points are.
    scored = [f for f in flags if f.criticality == "scored"]
    _heading(doc, f"3. Scored gaps ({len(scored)})")
    if scored:
        _table(
            doc,
            ["What's needed", "Owner", "Where", "Why", "If skipped"],
            [
                [
                    f.tag,
                    f.owner or "Unassigned",
                    f.section_title or f.section_id,
                    f.why_required or "Requested by the RFP.",
                    f.if_skipped or "Loses points.",
                ]
                for f in scored
            ],
        )
    else:
        doc.add_paragraph("None.")
    # "optional" flags are deliberately absent: they are removed from the manuscript,
    # so listing them here would re-introduce the noise that removal exists to clear.

    # 4 — Scorecard: the audit trail for the number in section 1.
    _heading(doc, "4. Scorecard")
    if scores:
        _table(
            doc,
            ["Criterion", "Weight", "Score", "What would lose points"],
            [
                [
                    s.criterion,
                    "unpublished" if s.weight is None else f"{s.weight:g}",
                    f"{s.score}/5",
                    "",
                ]
                for s in scores
            ],
        )
    else:
        doc.add_paragraph("No scored criteria were available for this RFP.")

    # 5 — What the reviewer changed.
    _heading(doc, "5. What the reviewer changed")
    if changes:
        for line in changes:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No changes recorded.")

    # 6 — Unverified claims: text kept, evidence not found, human confirmation required.
    _heading(doc, "6. Unverified claims")
    if unverified_claims:
        _note(
            doc,
            "Kept in the draft. The knowledge base could not confirm these — confirm "
            "before submitting.",
        )
        for line in unverified_claims:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("None.")

    # 7 — Convergence report: where the pipeline gave up, which is the tuning signal.
    _heading(doc, "7. What the reviewer could not fix")
    if unfixed:
        for line in unfixed:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("Nothing outstanding.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
