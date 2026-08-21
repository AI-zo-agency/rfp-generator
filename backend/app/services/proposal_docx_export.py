"""Build a Word (.docx) manuscript matching in-app preview structure."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models.proposal import ProposalDraft
from app.services.proposal_manuscript import build_manuscript_structured

_ORANGE = RGBColor(0xC2, 0x41, 0x0C)
_BODY = RGBColor(0x4B, 0x55, 0x63)
_MUTED = RGBColor(0x64, 0x74, 0x8B)


class ProposalDocxExportError(Exception):
    def __init__(self, message: str, *, status_code: int = 400) -> None:
        super().__init__(message)
        self.status_code = status_code


def _sanitize_filename(title: str) -> str:
    cleaned = re.sub(r'[^\w\s\-–—()]+', "", title or "").strip()
    base = (cleaned[:100] or "Proposal").strip()
    return f"{base} — Proposal.docx"


def _paragraph_shading(paragraph, fill_hex: str = "EFF6FF") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def _add_rich_paragraph(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    style: str | None = None,
    space_after_pt: float = 10,
) -> None:
    p = doc.add_paragraph(style=style)
    fmt = p.paragraph_format
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.25
    parts = re.split(r"(\*\*[^*]+\*\*)", text or "")
    for segment in parts:
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**"):
            run = p.add_run(segment[2:-2])
            run.bold = True
            run.font.color.rgb = _BODY
        else:
            run = p.add_run(segment)
            run.font.color.rgb = _BODY
            if italic:
                run.italic = True
    if not p.runs:
        p.add_run("")


def _add_designer_note(doc: Document, note: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(14)
    fmt.space_after = Pt(14)
    fmt.left_indent = Inches(0.12)
    fmt.right_indent = Inches(0.08)
    _paragraph_shading(p, "EFF6FF")

    label = p.add_run("DESIGNER NOTE\n")
    label.bold = True
    label.font.size = Pt(8)
    label.font.color.rgb = _MUTED

    body = p.add_run(note.strip())
    body.italic = True
    body.font.size = Pt(10.5)
    body.font.color.rgb = _MUTED


def _add_section_heading(doc: Document, title: str) -> None:
    h = doc.add_heading(title.strip(), level=1)
    fmt = h.paragraph_format
    fmt.space_before = Pt(18)
    fmt.space_after = Pt(8)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        run.font.size = Pt(14)


def _add_subheading(doc: Document, text: str, *, level: int) -> None:
    doc_level = min(max(level, 2), 3)
    h = doc.add_heading(text.strip(), level=doc_level)
    fmt = h.paragraph_format
    fmt.space_before = Pt(14)
    fmt.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = _ORANGE if doc_level >= 2 else RGBColor(0x11, 0x18, 0x27)
        if doc_level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)


_PAGE_INNER_INCHES = 6.5  # 8.5in letter minus 1in left/right margins
_DXA_PER_INCH = 1440
_WIDE_TABLE_COLS = 5
_FEE_HEADER_MARKERS = frozenset(
    {
        "phase",
        "deliverable",
        "amount",
        "item",
        "detail",
        "description",
        "extended",
        "hours",
        "rate",
        "cost",
        "total",
    }
)


def _looks_like_fee_or_form_table(headers: list[str]) -> bool:
    keys = {(h or "").strip().casefold() for h in headers}
    if "phase" in keys and ("amount" in keys or "deliverable" in keys):
        return True
    if "item" in keys and "detail" in keys and len(headers) <= 3:
        return True
    return len(keys & _FEE_HEADER_MARKERS) >= 2




def _set_table_full_width(table, col_count: int) -> None:
    """Stop WPS/Word autofit from collapsing columns to one-character width."""
    table.autofit = False
    try:
        table.allow_autofit = False
    except (AttributeError, ValueError):
        pass
    total = int(_PAGE_INNER_INCHES * _DXA_PER_INCH)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    def _set_child(tag: str, attrs: dict[str, str]) -> None:
        el = tbl_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tbl_pr.append(el)
        for key, value in attrs.items():
            el.set(qn(key), value)

    _set_child("w:tblW", {"w:w": str(total), "w:type": "dxa"})
    _set_child("w:tblLayout", {"w:type": "fixed"})
    _set_child("w:jc", {"w:val": "left"})

    if col_count <= 2:
        widths = [int(total * 0.34), total - int(total * 0.34)][:col_count]
    else:
        base = total // col_count
        widths = [base] * col_count
        widths[-1] += total - base * col_count

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
    else:
        grid = OxmlElement("w:tblGrid")
        tbl_pr.addnext(grid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx] if idx < len(widths) else widths[-1]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge, dxa in (("top", "60"), ("left", "80"), ("bottom", "60"), ("right", "80")):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), dxa)
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)


def _shape_table_for_word(
    headers: list[str],
    rows: list[list[str]],
) -> tuple[list[str], list[list[str]]]:
    """Wide one-row grids become a 2-column Item/Detail table so Word stays readable.

    Never reshape fee/phase tables — Word must keep Phase | Deliverable | Amount.
    """
    cols = max(len(headers), 1)
    hdr = (headers + [""] * cols)[:cols]
    data = [(list(row) + [""] * cols)[:cols] for row in rows]

    # Header-only fragments (blank lines split a markdown table) — do not turn
    # "Phase / Deliverable / Amount" into empty Item/Detail rows.
    if not data:
        return hdr, data

    if _looks_like_fee_or_form_table(hdr):
        return hdr, data

    if cols < _WIDE_TABLE_COLS:
        return hdr, data

    if len(data) <= 1:
        values = data[0] if data else [""] * cols
        # All-empty values → leave as-is (broken parse); reshaping makes it worse.
        if not any((v or "").strip() for v in values):
            return hdr, data
        return ["Item", "Detail"], [
            [hdr[i], values[i] if i < len(values) else ""] for i in range(cols)
        ]
    stacked: list[list[str]] = []
    for row in data:
        label = (row[0] or "").strip()
        for i, heading in enumerate(hdr):
            if i == 0:
                continue
            value = row[i] if i < len(row) else ""
            key = f"{label} — {heading}".strip(" —") if label else heading
            stacked.append([key, value])
    return ["Item", "Detail"], stacked or data


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    hdr, data = _shape_table_for_word(headers, rows)
    cols = max(len(hdr), 1)
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.style = "Table Grid"
    for c_idx, header in enumerate(hdr):
        cell = table.rows[0].cells[c_idx]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    _set_table_full_width(table, cols)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def build_proposal_docx_bytes(*, doc_title: str, draft: ProposalDraft) -> bytes:
    sections = build_manuscript_structured(draft.sections)
    if not sections:
        raise ProposalDocxExportError(
            "No proposal content to export. Generate sections first.",
            status_code=400,
        )

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    title = doc.add_heading(doc_title.strip(), level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    title.paragraph_format.space_after = Pt(16)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = _BODY

    for section in sections:
        _add_section_heading(doc, section.get("title") or "Untitled")

        for part in section.get("parts") or []:
            ptype = part.get("type")
            if ptype == "heading":
                text = (part.get("text") or "").strip()
                if text:
                    level = int(part.get("level") or 2)
                    mapped = min(max(level + 1, 2), 3)
                    _add_subheading(doc, text, level=mapped)
                continue
            if ptype == "designer_note":
                note = (part.get("text") or "").strip()
                if note:
                    _add_designer_note(doc, note)
                continue
            if ptype == "table":
                headers = part.get("headers") or []
                rows = part.get("rows") or []
                if headers:
                    _add_table(doc, headers, rows)
                continue
            if ptype == "list":
                items = part.get("items") or []
                ordered = bool(part.get("ordered"))
                style = "List Number" if ordered else "List Bullet"
                for item in items:
                    _add_rich_paragraph(
                        doc,
                        (item or "").strip(),
                        style=style,
                        space_after_pt=4,
                    )
                gap = doc.add_paragraph()
                gap.paragraph_format.space_after = Pt(6)
                continue
            text = (part.get("text") or "").strip()
            if text:
                for chunk in re.split(r"\n{2,}", text):
                    chunk = chunk.strip()
                    if chunk:
                        _add_rich_paragraph(doc, chunk)

        doc.add_paragraph()  # section spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_proposal_docx_filename(*, rfp_title: str) -> str:
    return _sanitize_filename(rfp_title)
