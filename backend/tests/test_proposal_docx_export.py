"""Word export must keep tables readable (no one-character column collapse)."""

from __future__ import annotations

import io
import unittest

from docx import Document
from docx.oxml.ns import qn

from app.models.proposal import ProposalDraft, ProposalSection
from app.services.proposal_manuscript import parse_markdown_parts, repair_flattened_markdown_tables
from app.services.proposal_docx_export import (
    _shape_table_for_word,
    build_proposal_docx_bytes,
)


class WordTableLayoutTests(unittest.TestCase):
    def test_wide_one_row_table_becomes_item_detail(self) -> None:
        headers = [
            "Years in Operation",
            "Business Registrations",
            "Legal Structure",
            "Certifications",
            "Grant Experience",
            "Invoicing Readiness",
        ]
        rows = [
            [
                "13 years",
                "Oregon, Washington, Texas, Colorado, and California",
                "S-Corp/LLC, sole owner Sonja Anderson",
                "WBENC Women's Business Enterprise",
                "Public health / behavioral health campaigns",
                "Quarterly invoicing aligned to grant spend-down",
            ]
        ]
        hdr, data = _shape_table_for_word(headers, rows)
        self.assertEqual(hdr, ["Item", "Detail"])
        self.assertEqual(len(data), 6)
        self.assertEqual(data[1][0], "Business Registrations")
        self.assertIn("Oregon", data[1][1])

    def test_fee_phase_table_stays_three_columns(self) -> None:
        hdr, data = _shape_table_for_word(
            ["Phase", "Deliverable", "Amount"],
            [
                ["Phase 1: Discovery", "Stakeholder interviews", "$6,639.60"],
                ["Phase 2: Strategy", "Campaign platform", "$4,742.57"],
            ],
        )
        self.assertEqual(hdr, ["Phase", "Deliverable", "Amount"])
        self.assertEqual(len(data), 2)
        self.assertEqual(data[0][2], "$6,639.60")

    def test_pipe_in_cell_text_does_not_explode_columns(self) -> None:
        """A literal | inside deliverable text must not create extra columns."""
        md = (
            "| Phase | Deliverable | Amount |\n"
            "| --- | --- | ---: |\n"
            "| Phase 1 | Stakeholder interviews (RFQ Section 3, Scope of Work) | $6,639.60 |\n"
            "| Phase 2 | Campaign platform | $4,742.57 |\n"
        )
        parts = parse_markdown_parts(md)
        self.assertEqual(len(parts), 1)
        tbl = parts[0]
        self.assertEqual(tbl["type"], "table")
        self.assertEqual(len(tbl["headers"]), 3)
        for row in tbl["rows"]:
            self.assertEqual(len(row), 3, f"Row has wrong column count: {row}")
        self.assertIn("$6,639.60", tbl["rows"][0][2])

    def test_flattened_one_line_table_is_repaired_in_markdown(self) -> None:
        raw = (
            "## Fee Detail by Phase\n\n"
            "| Phase | Deliverable | Amount | | --- | --- | ---: | "
            "| Phase 1: Discovery | Stakeholder interviews | $6,639.60 | "
            "| Phase 2: Strategy | Campaign platform | $4,742.57 |"
        )
        repaired = repair_flattened_markdown_tables(raw)
        self.assertIn("| Phase | Deliverable | Amount |\n", repaired)
        self.assertIn("| --- | --- | --- |\n", repaired)
        self.assertIn("| Phase 1: Discovery | Stakeholder interviews | $6,639.60 |\n", repaired)
        self.assertNotIn("| Amount | | --- |", repaired)
        parts = parse_markdown_parts(repaired)
        tables = [p for p in parts if p["type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0]["headers"], ["Phase", "Deliverable", "Amount"])
        self.assertEqual(len(tables[0]["rows"]), 2)
        self.assertEqual(tables[0]["rows"][0][2], "$6,639.60")

    def test_flattened_table_keeps_empty_total_cell(self) -> None:
        raw = (
            "| Phase | Deliverable | Amount | | --- | --- | ---: | "
            "| Phase 1 | Discovery | $1,000 | | **Total** | | **$1,000** |"
        )
        repaired = repair_flattened_markdown_tables(raw)
        parts = parse_markdown_parts(repaired)
        table = next(p for p in parts if p["type"] == "table")
        self.assertEqual(len(table["headers"]), 3)
        last = table["rows"][-1]
        self.assertEqual(len(last), 3)
        self.assertIn("Total", last[0])
        self.assertIn("$1,000", last[2])

    def test_manual_fill_with_pipe_is_not_a_table_row(self) -> None:
        raw = (
            "| Phase | Deliverable | Amount |\n"
            "| --- | --- | ---: |\n"
            "| Phase 1 | Discovery | $1,000 |\n"
            "| **Total** | | **$1,000** |\n\n"
            "[MANUAL FILL: Sonja — leftover note | still a flag]\n"
        )
        repaired = repair_flattened_markdown_tables(raw)
        self.assertNotIn("| [MANUAL FILL", repaired)
        self.assertIn("[MANUAL FILL:", repaired)
        self.assertNotIn("leftover note | still", repaired)
        parts = parse_markdown_parts(repaired)
        tables = [p for p in parts if p["type"] == "table"]
        self.assertEqual(len(tables), 1)
        last = tables[0]["rows"][-1]
        self.assertIn("Total", last[0])
        self.assertNotIn("MANUAL FILL", last[0])

    def test_verify_cells_inside_addenda_table_stay_as_table(self) -> None:
        raw = (
            "| | **Addendum Number** | **Date Issued** | **Acknowledged** | **Incorporated** |\n"
            "| --- | --- | --- | --- | --- |\n"
            "| | [VERIFY: Sonja confirm addendum number and date from Bonfire portal] | "
            "[VERIFY: Sonja confirm addendum number and date from Bonfire portal] | Yes | Yes |\n"
        )
        repaired = repair_flattened_markdown_tables(raw)
        self.assertIn("| Yes | Yes |", repaired)
        self.assertNotIn("] — Yes — Yes", repaired)
        parts = parse_markdown_parts(repaired)
        tables = [p for p in parts if p["type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(len(tables[0]["headers"]), 5)
        self.assertEqual(len(tables[0]["rows"]), 1)
        self.assertEqual(tables[0]["rows"][0][-2], "Yes")
        self.assertEqual(tables[0]["rows"][0][-1], "Yes")

    def test_normal_table_with_empty_cell_is_not_split(self) -> None:
        raw = (
            "| Phase | Notes | Amount |\n"
            "| --- | --- | ---: |\n"
            "| Phase 1 | | $5,000 |\n"
            "| **Total** | | **$5,000** |\n"
        )
        self.assertEqual(repair_flattened_markdown_tables(raw), raw)

    def test_blank_lines_inside_fee_table_still_export_as_one_grid(self) -> None:
        draft = ProposalDraft(
            rfpId="rfp-1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="price",
                    title="Price — All-Inclusive Cost Proposal",
                    content=(
                        "## Fee Detail by Phase\n\n"
                        "| Phase | Deliverable | Amount |\n"
                        "\n"
                        "| --- | --- | ---: |\n"
                        "\n"
                        "| Phase 1: Discovery & Audience Research | "
                        "Stakeholder interviews with MCECS | $6,639.60 |\n"
                        "\n"
                        "| Phase 2: Strategy | Campaign platform | $4,742.57 |\n"
                    ),
                    status="generated",
                )
            ],
        )
        blob = build_proposal_docx_bytes(doc_title="PSU", draft=draft)
        doc = Document(io.BytesIO(blob))
        self.assertTrue(doc.tables)
        table = doc.tables[0]
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(len(table.rows), 3)  # header + 2 data
        texts = [[cell.text for cell in row.cells] for row in table.rows]
        self.assertEqual(texts[0], ["Phase", "Deliverable", "Amount"])
        self.assertIn("Phase 1", texts[1][0])
        self.assertIn("$6,639.60", texts[1][2])
        self.assertNotIn("---", texts[1][0])
        # Must not collapse into empty Item/Detail rows.
        self.assertFalse(
            any(row == ["Item", "Detail"] for row in texts)
        )

    def test_header_only_wide_fragment_does_not_become_empty_item_detail(self) -> None:
        hdr, data = _shape_table_for_word(
            ["Phase", "Deliverable", "Amount", "Hours", "Rate", "Notes"],
            [],
        )
        self.assertEqual(hdr[0], "Phase")
        self.assertEqual(data, [])


    def test_exported_tables_have_fixed_page_width(self) -> None:
        draft = ProposalDraft(
            rfpId="rfp-1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="financial",
                    title="Financial Stability",
                    content=(
                        "## Financial Stability\n\n"
                        "| Years in Operation | Business Registrations | Legal Structure | "
                        "Certifications | Grant Experience |\n"
                        "| --- | --- | --- | --- | --- |\n"
                        "| 13 years | Oregon, Washington, Texas, Colorado, and California | "
                        "S-Corp/LLC | WBENC | Public health campaigns |\n"
                    ),
                    status="generated",
                )
            ],
        )
        blob = build_proposal_docx_bytes(doc_title="Calvert County", draft=draft)
        doc = Document(io.BytesIO(blob))
        self.assertTrue(doc.tables)
        table = doc.tables[0]
        self.assertEqual(len(table.columns), 2)
        texts = [cell.text for row in table.rows for cell in row.cells]
        self.assertTrue(any("Oregon" in t for t in texts))
        self.assertTrue(any("Business Registrations" in t for t in texts))
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        self.assertIsNotNone(tbl_w)
        self.assertEqual(tbl_w.get(qn("w:type")), "dxa")
        self.assertGreater(int(tbl_w.get(qn("w:w"))), 8000)
        self.assertEqual(layout.get(qn("w:type")), "fixed")


if __name__ == "__main__":
    unittest.main()
