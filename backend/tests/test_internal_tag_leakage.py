"""No internal tag may survive to a client-facing export.

Observed: [PRICING FLAG: ... outside guide band ...] reached build_manuscript_blocks,
and DOCX export renders [DESIGNER NOTE] as a styled block.
"""

from __future__ import annotations

import re
import unittest

from app.models.proposal import ProposalDraft, ProposalSection
from app.services.evidence_trust.flags import (
    flag_claim_mismatch,
    flag_confirm,
    flag_provenance,
)
from app.services.proposal_docx_export import build_proposal_docx_bytes
from app.services.proposal_manuscript import (
    build_manuscript_blocks,
    build_manuscript_structured,
    plain_text_for_export,
)
from app.services.proposal_presubmit_review import _PLACEHOLDER_RE
from app.services.proposal_t1_validators import scan_internal_note_leaks

# Bare [FLAG: ...] tags are emitted for real by app/services/evidence_trust/flags.py
# and reach section.content via claim_validator.validate_and_flag_section (from
# proposal_fulfill_fabrication_guard.repair_fabricated_qualifications and
# proposal_chat_ops.py) — the same path this module guards. Built from the real
# emitters so these tests break if the emitted tag shape ever changes.
BARE_FLAG_TAGS = [
    flag_confirm("Acme Corp"),
    flag_claim_mismatch("Acme Corp", "led the rebrand", "media buying"),
    flag_provenance("blog.example.com", "unverified source"),
]

INTERNAL_TAGS = [
    "[PRICING FLAG: PM ratio 12% outside 5-8% guide]",
    "[MANUAL FILL: Sonja — confirm commission rate]",
    "[FLAG FOR SONJA: Recovery Network of Oregon]",
    "[DESIGNER NOTE: insert 04_Bio_SonjaAnderson.pdf]",
    "[VERIFY: years in operation]",
    *BARE_FLAG_TAGS,
]

# "FLAG:" covers both bare [FLAG: ...] and [PRICING FLAG: ...].
_MARKERS = (
    "PRICING FLAG",
    "MANUAL FILL",
    "FLAG FOR",
    "DESIGNER NOTE",
    "VERIFY:",
    "FLAG:",
)


class PlaceholderScannerTests(unittest.TestCase):
    def test_every_internal_tag_is_detected(self) -> None:
        for tag in INTERNAL_TAGS:
            with self.subTest(tag=tag):
                self.assertTrue(_PLACEHOLDER_RE.findall(tag), f"{tag} not detected")

    def test_ordinary_brackets_are_not_flagged(self) -> None:
        self.assertFalse(_PLACEHOLDER_RE.findall("[1] See appendix A [2]"))


class ExportScrubTests(unittest.TestCase):
    def test_no_internal_tag_survives_plain_text_export(self) -> None:
        body = "We deliver the plan.\n\n" + "\n\n".join(INTERNAL_TAGS)
        out = plain_text_for_export(body)
        for marker in _MARKERS:
            with self.subTest(marker=marker):
                self.assertNotIn(marker, out)

    def test_real_content_survives_the_scrub(self) -> None:
        out = plain_text_for_export("We deliver the plan.\n\n[PRICING FLAG: x]")
        self.assertIn("We deliver the plan.", out)


def _draft_with_all_tags() -> ProposalDraft:
    body = "We deliver the plan on time.\n\n" + "\n\n".join(INTERNAL_TAGS)
    section = ProposalSection(
        id="section-4-approach",
        title="Approach",
        content=body,
        status="generated",
    )
    return ProposalDraft(
        rfpId="r1",
        updatedAt="2026-01-01T00:00:00Z",
        sections=[section],
    )


class RealExportPathTests(unittest.TestCase):
    """Drive the actual export functions end to end, not just the helper regexes."""

    def test_build_manuscript_blocks_scrubs_every_tag(self) -> None:
        draft = _draft_with_all_tags()
        blocks = build_manuscript_blocks(draft.sections)
        self.assertTrue(blocks)
        combined = "\n".join(body for _title, body in blocks)
        for marker in _MARKERS:
            with self.subTest(marker=marker):
                self.assertNotIn(marker, combined)
        self.assertIn("We deliver the plan on time.", combined)

    def test_build_manuscript_structured_scrubs_every_tag(self) -> None:
        draft = _draft_with_all_tags()
        structured = build_manuscript_structured(draft.sections)
        self.assertTrue(structured)
        combined = "\n".join(
            part.get("text", "")
            for sec in structured
            for part in sec.get("parts", [])
        )
        for marker in _MARKERS:
            with self.subTest(marker=marker):
                self.assertNotIn(marker, combined)
        # No part should ever be typed as a designer_note — it must never
        # reach the exported document at all, styled or otherwise.
        part_types = {
            part.get("type")
            for sec in structured
            for part in sec.get("parts", [])
        }
        self.assertNotIn("designer_note", part_types)

    def test_docx_export_bytes_contain_no_internal_tag(self) -> None:
        """Drive the real DOCX export path — regressions here shipped [DESIGNER
        NOTE] as a rendered styled block straight into the submitted document."""
        import io

        from docx import Document

        draft = _draft_with_all_tags()
        raw = build_proposal_docx_bytes(doc_title="Test Proposal", draft=draft)
        doc = Document(io.BytesIO(raw))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for cell_text in (
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ):
            full_text += "\n" + cell_text
        for marker in _MARKERS:
            with self.subTest(marker=marker):
                self.assertNotIn(marker, full_text)
        self.assertIn("We deliver the plan on time.", full_text)


class T1WhitelistTests(unittest.TestCase):
    """Narrowing _ALLOWED_BRACKET_TAG_RE to VERIFY must not create new blockers
    for legitimate MANUAL FILL / DESIGNER NOTE tags that survive authoring."""

    def test_manual_fill_still_not_a_note_leak(self) -> None:
        draft = _draft_with_all_tags()
        findings = scan_internal_note_leaks(draft)
        codes = {f["code"] for f in findings}
        self.assertNotIn("t1.note_leak.manual_fill", codes)
        for f in findings:
            self.assertNotIn("MANUAL FILL", f["message"].upper())

    def test_designer_note_still_not_a_note_leak(self) -> None:
        draft = ProposalDraft(
            rfpId="r1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="s1",
                    title="Team",
                    content="Layout: [DESIGNER NOTE: place map on facing page].",
                    status="generated",
                )
            ],
        )
        findings = scan_internal_note_leaks(draft)
        self.assertEqual(findings, [])

    def test_pricing_flag_is_a_note_leak(self) -> None:
        draft = ProposalDraft(
            rfpId="r1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="s1",
                    title="Budget",
                    content="Fees $3,500. [PRICING FLAG: line l1 unbound — outside guide band]",
                    status="generated",
                )
            ],
        )
        findings = scan_internal_note_leaks(draft)
        self.assertTrue(
            any(f["code"] == "t1.note_leak.pricing_flag" for f in findings)
        )
        self.assertTrue(any(f["blocker"] for f in findings))

    def test_bare_flag_is_not_a_t1_blocker(self) -> None:
        """Deliberate: evidence_trust emits bare [FLAG: ...] during authoring, so
        making it a T1 *blocker* would halt legitimate proposals. It is surfaced
        by the presubmit placeholder scanner and removed by the export scrub
        instead — see the decision table in task-5-report.md."""
        draft = ProposalDraft(
            rfpId="r1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="s1",
                    title="Our Work",
                    content=f"We served {flag_confirm('Acme Corp')} last year.",
                    status="generated",
                )
            ],
        )
        self.assertEqual(scan_internal_note_leaks(draft), [])
        # ...but the presubmit scanner must still see it.
        self.assertTrue(_PLACEHOLDER_RE.findall(flag_confirm("Acme Corp")))


class BareFlagTagExportTests(unittest.TestCase):
    """Regression: evidence_trust/flags.py emits bare [FLAG: ...] (no 'FOR'), which
    the first version of the export scrub missed entirely — it only matched
    'FLAG FOR'. Verified leaking into both plain text and real DOCX bytes."""

    def _draft(self) -> ProposalDraft:
        body = "We served three public universities.\n\n" + "\n\n".join(
            BARE_FLAG_TAGS
        )
        return ProposalDraft(
            rfpId="r1",
            updatedAt="2026-01-01T00:00:00Z",
            sections=[
                ProposalSection(
                    id="section-3-our-work",
                    title="Our Work",
                    content=body,
                    status="generated",
                )
            ],
        )

    def test_bare_flag_does_not_survive_plain_text_export(self) -> None:
        for tag in BARE_FLAG_TAGS:
            with self.subTest(tag=tag):
                out = plain_text_for_export(f"We served Acme Corp.\n\n{tag}")
                self.assertNotIn("FLAG", out)
                self.assertIn("We served Acme Corp.", out)

    def test_bare_flag_does_not_survive_build_manuscript_structured(self) -> None:
        structured = build_manuscript_structured(self._draft().sections)
        combined = "\n".join(
            part.get("text", "")
            for sec in structured
            for part in sec.get("parts", [])
        )
        self.assertNotIn("FLAG", combined)
        self.assertIn("We served three public universities.", combined)

    def test_bare_flag_does_not_survive_real_docx_bytes(self) -> None:
        import io

        from docx import Document

        raw = build_proposal_docx_bytes(doc_title="Test", draft=self._draft())
        doc = Document(io.BytesIO(raw))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        self.assertNotIn("FLAG", full_text)
        self.assertIn("We served three public universities.", full_text)


class ScrubDoesNotOverReachTests(unittest.TestCase):
    """The shared tag regex is word-boundary anchored — ordinary prose that merely
    starts with the same letters must not be eaten by the export scrub."""

    def test_ordinary_bracketed_citations_survive(self) -> None:
        out = plain_text_for_export("[1] See appendix A [2]")
        self.assertIn("[1]", out)
        self.assertIn("[2]", out)
        self.assertIn("See appendix A", out)

    def test_words_merely_starting_with_a_tag_name_survive(self) -> None:
        for text in (
            "Our [FLAGSHIP PROGRAM] served the county.",
            "See [VERIFICATION SUMMARY] for details.",
        ):
            with self.subTest(text=text):
                self.assertEqual(plain_text_for_export(text), text)


if __name__ == "__main__":
    unittest.main()
