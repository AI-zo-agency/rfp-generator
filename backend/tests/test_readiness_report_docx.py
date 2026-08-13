"""Submission Readiness Report — an internal document, never the evaluator's.

Rows are written for the person doing the work: a blocker names the owner, the clause,
and the consequence. "Missing bid bond" is not an actionable row.
"""

from __future__ import annotations

import io

from docx import Document

from app.models.proposal import ManualFillFlag
from app.services.proposal_readiness import CriterionScore, compute_readiness
from app.services.proposal_readiness_report import (
    build_readiness_report_docx_bytes,
    build_readiness_report_filename,
)


def _text_of(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _blocker() -> ManualFillFlag:
    return ManualFillFlag(
        sectionId="s7",
        sectionTitle="Submission Forms",
        kind="manual_fill",
        tag="Bid bond",
        owner="Sonja",
        criticality="disqualifying",
        rfpEvidence="Bids submitted without a bid bond will be rejected as non-responsive.",
        whyRequired="The RFP rejects any bid with no bond.",
        ifSkipped="Bid rejected unopened. Not scored at all.",
    )


def _scored() -> ManualFillFlag:
    return ManualFillFlag(
        sectionId="s2",
        sectionTitle="Experience",
        kind="manual_fill",
        tag="Third client reference",
        owner="Ella",
        criticality="scored",
        whyRequired="The RFP asks for three references; two are present.",
        ifSkipped="Loses points on Experience.",
    )


def _optional() -> ManualFillFlag:
    return ManualFillFlag(
        sectionId="s3",
        sectionTitle="Team",
        kind="manual_fill",
        tag="Team photographs",
        criticality="optional",
    )


def _report(flags, scores=None, unresolved=0, **kw) -> str:
    scores = scores or [CriterionScore("s1", "Approach", 4, 1.0)]
    readiness = compute_readiness(scores=scores, flags=flags, unresolved=unresolved)
    return _text_of(
        build_readiness_report_docx_bytes(
            rfp_title="City Marketing RFP",
            readiness=readiness,
            flags=flags,
            scores=scores,
            **kw,
        )
    )


class TestVerdict:
    def test_score_and_verdict_are_present(self):
        text = _report([_blocker()])
        assert "%" in text
        assert "disqualifying" in text.casefold()

    def test_confidence_is_stated(self):
        text = _report(
            [], scores=[CriterionScore("s1", "A", 4, None) for _ in range(4)]
        )
        assert "confidence" in text.casefold()
        assert "weights unpublished" in text.casefold()


class TestBlockers:
    def test_blocker_row_is_actionable_not_a_label(self):
        text = _report([_blocker()])
        assert "Bid bond" in text
        assert "Sonja" in text  # owner
        assert "Submission Forms" in text  # where
        assert "rejected as non-responsive" in text  # the quoted clause
        assert "Bid rejected unopened" in text  # if skipped

    def test_blockers_appear_before_scored_gaps(self):
        text = _report([_scored(), _blocker()])
        assert text.index("Bid bond") < text.index("Third client reference")


class TestScoredGaps:
    def test_scored_gap_is_listed(self):
        text = _report([_scored()])
        assert "Third client reference" in text
        assert "Loses points" in text


class TestOptionalIsAbsent:
    def test_optional_flags_never_reach_the_report(self):
        """Noise stays gone — optional items are removed from the manuscript."""
        text = _report([_optional(), _scored()])
        assert "Team photographs" not in text
        assert "Third client reference" in text


class TestOtherSections:
    def test_scorecard_lists_criteria_and_weights(self):
        text = _report(
            [],
            scores=[CriterionScore("s1", "Technical Approach", 3, 40.0)],
        )
        assert "Technical Approach" in text

    def test_unverified_claims_are_listed(self):
        text = _report([], unresolved=2, unverified_claims=["Acme saved 30%"])
        assert "Acme saved 30%" in text

    def test_convergence_report_is_listed(self):
        text = _report([], unfixed=["Section 3 repetition survived 3 rounds"])
        assert "survived 3 rounds" in text

    def test_changes_made_are_summarised(self):
        text = _report([], changes=["Swapped case study in Section 4"])
        assert "Swapped case study" in text


class TestDegenerate:
    def test_clean_draft_still_produces_a_valid_document(self):
        text = _report([])
        assert "%" in text

    def test_filename_is_stable_and_safe(self):
        name = build_readiness_report_filename(rfp_title="City / Marketing: RFP")
        assert name.endswith(".docx")
        assert "/" not in name
